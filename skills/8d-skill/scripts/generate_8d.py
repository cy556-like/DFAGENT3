#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
8D 问题解决报告生成器
=====================
接收产品名、缺陷描述、客户名等参数，输出 8D 报告 .xlsx（单 Sheet）和 .docx 文件。

依赖：
    - openpyxl（生成 Excel）
    - python-docx（生成 Word）

未安装时自动 pip install。

用法：
    python3 generate_8d.py \\
        --product "前保险杠总成（涂装件）" \\
        --defect "漆面颗粒/杂质" \\
        --customer "XX汽车有限公司" \\
        --defect-rate "500PPM" \\
        --batch-size "15" \\
        --template paint-defect \\
        --output-dir ~/Desktop

⚠️ 行业常识基准（详见 SKILL.md 第十章）：
    - defect-rate 默认示例用 500 PPM (0.05%)，安全件用 50 PPM；
      严禁使用 3% / 5% / 8% / 11.5% 等灾难级数字作为示例。
    - batch-size 字段语义为「本次 8D 分析的客户投诉/退货样本件数」，
      典型 1–30 件；不是生产批量，严禁使用 500 / 2000 等数字作为示例。
"""

import argparse
import json
import os
import subprocess
import sys
import datetime
from pathlib import Path

# ============================================================
# 依赖检查与自动安装
# ============================================================

REQUIRED_PACKAGES = {
    "openpyxl": "openpyxl",
    "docx": "python-docx",  # 注意：导入名是 docx，包名是 python-docx
}


def ensure_packages():
    """检查并自动安装所需的 Python 包。"""
    import importlib

    for import_name, pip_name in REQUIRED_PACKAGES.items():
        try:
            importlib.import_module(import_name)
        except ImportError:
            print(f"[INFO] {pip_name} 未安装，正在自动安装...")
            try:
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", pip_name, "--quiet"]
                )
                print(f"[INFO] {pip_name} 安装完成。")
            except subprocess.CalledProcessError as e:
                print(f"[ERROR] {pip_name} 安装失败：{e}")
                print(f"        请手动执行：pip install {pip_name}")
                sys.exit(1)


ensure_packages()

# 重新导入确保可用
import openpyxl
from openpyxl.styles import (
    Alignment,
    Border,
    Side,
    PatternFill,
    Font,
)
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ============================================================
# 常量
# ============================================================

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent  # 8d-report/
TEMPLATES_DIR = SKILL_ROOT / "templates"

# 颜色
HEADER_FILL = "003366"  # 深蓝色表头
ALT_ROW_FILL = "D6E4F0"  # 交替行浅蓝
ROOT_CAUSE_FILL = "FFF2CC"  # 根本原因行黄色高亮
SUBHEADER_FILL = "4472C4"  # 次级表头
YELLOW_SECTION_FILL = "FFF8E1"  # 章节标题淡黄色底

# 字体
FONT_NAME = "微软雅黑"
FONT_NAME_EN = "Microsoft YaHei"


# ============================================================
# 工具函数
# ============================================================

def normalize_defect_rate(value):
    """统一不良率表达：接受 '500PPM' / '500ppm' / '0.05%' / '0.05' 等格式，
    返回 '500 PPM (0.05%)' 这种双格式字符串，便于客户/内部双向理解。

    遵循 SKILL.md 第十章「行业常识基准」：>0.5% 视为严重事故，函数会标注警告但不会拒绝。
    """
    if not value or value == "____" or not isinstance(value, str):
        return value if value else "____"
    s = value.strip()
    # 已经是双格式（含 PPM 和 %）
    if "PPM" in s.upper() and "%" in s:
        return s
    # 提取数字
    import re
    m = re.search(r"(\d+(?:\.\d+)?)", s)
    if not m:
        return s  # 无法解析，原样返回
    num = float(m.group(1))
    upper = s.upper()
    if "PPM" in upper:
        # PPM → %
        percent = num / 10000.0
        return f"{int(num) if num.is_integer() else num} PPM ({percent:.4f}%)"
    elif "%" in s:
        # % → PPM
        ppm = num * 10000
        return f"{int(ppm) if ppm.is_integer() else ppm} PPM ({num}%)"
    else:
        # 纯数字，按行业常识判断：≤1 视为百分比，>1 视为 PPM
        if num <= 1:
            ppm = num * 10000
            return f"{int(ppm) if ppm.is_integer() else ppm} PPM ({num}%)"
        else:
            percent = num / 10000.0
            return f"{int(num) if num.is_integer() else num} PPM ({percent:.4f}%)"


def replace_placeholders(text, context):
    """替换文本中的占位符。"""
    if not isinstance(text, str):
        return text
    return (
        text.replace("{defect_type}", context.get("defect", ""))
        .replace("{product_name}", context.get("product", ""))
        .replace("{customer}", context.get("customer", ""))
    )


def replace_placeholders_deep(obj, context):
    """递归替换字典/列表中的占位符。"""
    if isinstance(obj, str):
        return replace_placeholders(obj, context)
    if isinstance(obj, list):
        return [replace_placeholders_deep(item, context) for item in obj]
    if isinstance(obj, dict):
        return {k: replace_placeholders_deep(v, context) for k, v in obj.items()}
    return obj


def generate_8d_number():
    """生成 8D 编号：8D-YYYYMMDD-HHMMSS"""
    now = datetime.datetime.now()
    return now.strftime("8D-%Y%m%d-%H%M%S")


def safe_filename(name):
    """将字符串转换为安全的文件名（去除非法字符）。"""
    illegal_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|', '\n', '\r', '\t']
    for ch in illegal_chars:
        name = name.replace(ch, '_')
    # 截断过长的文件名
    if len(name) > 50:
        name = name[:50]
    return name


def load_template(template_slug, context):
    """加载模板 JSON 并替换占位符。"""
    template_path = TEMPLATES_DIR / template_slug / "template.json"
    if not template_path.exists():
        print(f"[WARN] 模板 {template_slug} 不存在，回退到 generic-defect")
        template_path = TEMPLATES_DIR / "generic-defect" / "template.json"

    with open(template_path, "r", encoding="utf-8") as f:
        template = json.load(f)

    # 递归替换占位符
    template = replace_placeholders_deep(template, context)
    return template


# ============================================================
# Excel 生成
# ============================================================

def get_thin_border():
    """获取细边框。"""
    side = Side(border_style="thin", color="000000")
    return Border(left=side, right=side, top=side, bottom=side)


def get_header_font():
    return Font(name=FONT_NAME, size=11, bold=True, color="FFFFFF")


def get_body_font():
    return Font(name=FONT_NAME, size=10, color="000000")


def get_header_fill():
    return PatternFill(start_color=HEADER_FILL, end_color=HEADER_FILL, fill_type="solid")


def get_alt_fill():
    return PatternFill(start_color=ALT_ROW_FILL, end_color=ALT_ROW_FILL, fill_type="solid")


def get_root_cause_fill():
    return PatternFill(start_color=ROOT_CAUSE_FILL, end_color=ROOT_CAUSE_FILL, fill_type="solid")


def get_subheader_fill():
    return PatternFill(start_color=SUBHEADER_FILL, end_color=SUBHEADER_FILL, fill_type="solid")


def get_subheader_font():
    return Font(name=FONT_NAME, size=10, bold=True, color="FFFFFF")


def get_yellow_section_fill():
    """获取淡黄色章节标题填充。"""
    return PatternFill(start_color=YELLOW_SECTION_FILL, end_color=YELLOW_SECTION_FILL, fill_type="solid")


def get_yellow_section_font():
    """获取黄色章节标题字体（深蓝字）。"""
    return Font(name=FONT_NAME, size=12, bold=True, color=HEADER_FILL)


def apply_yellow_section_style(cell):
    """应用黄色章节标题样式。"""
    cell.font = get_yellow_section_font()
    cell.fill = get_yellow_section_fill()
    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_header_style(cell):
    """应用表头样式。"""
    cell.font = get_header_font()
    cell.fill = get_header_fill()
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_body_style(cell, row_idx_in_data, is_root_cause=False):
    """应用正文样式（交替行 / 根因高亮）。"""
    cell.font = get_body_font()
    if is_root_cause:
        cell.fill = get_root_cause_fill()
    elif row_idx_in_data % 2 == 1:
        cell.fill = get_alt_fill()
    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_subheader_style(cell):
    """应用次级表头样式。"""
    cell.font = get_subheader_font()
    cell.fill = get_subheader_fill()
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def _display_width(s: str) -> int:
    """计算字符串的显示宽度（中文≈2，英文≈1），用于行高估算。"""
    w = 0
    for ch in str(s):
        if '\u4e00' <= ch <= '\u9fff' or '\u3000' <= ch <= '\u303f' or '\uff00' <= ch <= '\uffef':
            w += 2
        else:
            w += 1
    return w


def _calc_row_height(row_data, col_widths, line_height=22, min_height=25, max_height=409):
    """根据单元格内容和列宽精确计算行高，确保所有文字完全显示。

    核心逻辑：
    1. 遍历当前行每个单元格，计算该单元格文本在该列宽度下需要换几行
    2. 取所有单元格中最大的行数（即"最占行数的那一列"决定整行高度）
    3. 行高 = 最大行数 × 每行高度 + 上下边距

    CJK 字符宽度按英文2倍计算，显式换行符 \\n 也增加行数。

    Args:
        row_data: 该行各列的值列表
        col_widths: 各列宽度列表（Excel 列宽单位，约等于英文字符数）
        line_height: 每行高度（磅），默认22，确保10-11号字完整显示
        min_height: 最小行高
        max_height: 最大行高（Excel 最大409磅）

    Returns:
        int: 计算出的行高（磅）
    """
    max_lines = 1
    for col_idx, value in enumerate(row_data):
        text = str(value) if value is not None else ""
        if not text:
            continue
        # 获取该列宽度（默认10）
        col_w = col_widths[col_idx] if col_idx < len(col_widths) else 10
        # 实际文本区宽度：列宽减去单元格内边距（约4字符宽）
        # 必须留足边距，否则估算行数偏少导致文字被截断
        effective_width = max(col_w - 4, 3)

        # 按显式换行符分段计算
        segments = text.split('\n')
        total_lines = 0
        for segment in segments:
            if not segment.strip():
                total_lines += 1  # 空行也算一行
                continue
            seg_display_w = _display_width(segment)
            # 向上取整计算该段需要几行
            lines_needed = -(-seg_display_w // effective_width)  # 等价于 ceil(a/b)
            lines_needed = max(1, lines_needed)
            total_lines += lines_needed

        max_lines = max(max_lines, total_lines)

    # 行高 = 行数 × 每行高度 + 上下边距(8磅)
    calculated_height = max_lines * line_height + 8
    return max(min_height, min(max_height, calculated_height))


def _recalc_all_row_heights(ws, line_height=22, min_height=25, max_height=409):
    """重新计算整个工作表所有行的行高（基于实际单元格内容和列宽）。

    这是在所有内容写入完成、auto_fill 执行完毕后调用的最终修正步骤。
    确保每个单元格的文字都完整显示，行高由该行文字最多的列决定。

    处理逻辑：
    1. 遍历每一行
    2. 对每个有内容的单元格，根据文本长度和列宽计算所需行数
    3. 合并单元格的宽度 = 涉及所有列宽之和
    4. 取该行所有单元格中最大行数，计算行高
    """
    for row_idx in range(1, ws.max_row + 1):
        max_lines = 1
        has_content = False
        for col_idx in range(1, ws.max_column + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            if cell.value is None:
                continue
            text = str(cell.value)
            if not text.strip():
                continue
            has_content = True

            # 获取该列宽度
            col_letter = get_column_letter(col_idx)
            col_w = ws.column_dimensions[col_letter].width or 10

            # 如果是合并单元格，累加所有合并列的宽度
            effective_width = col_w
            for merge_range in ws.merged_cells.ranges:
                if cell.coordinate in merge_range:
                    total_w = 0
                    for c in range(merge_range.min_col, merge_range.max_col + 1):
                        cl = get_column_letter(c)
                        total_w += ws.column_dimensions[cl].width or 10
                    effective_width = total_w
                    break

            # 计算行数
            effective_width = max(effective_width - 4, 3)
            segments = text.split('\n')
            total_lines = 0
            for segment in segments:
                if not segment.strip():
                    total_lines += 1
                    continue
                seg_w = _display_width(segment)
                lines_needed = max(1, -(-seg_w // effective_width))
                total_lines += lines_needed
            max_lines = max(max_lines, max(1, total_lines))

        if has_content:
            calculated_height = max_lines * line_height + 8
            ws.row_dimensions[row_idx].height = max(min_height, min(max_height, calculated_height))


def set_column_widths(ws, widths):
    """设置列宽。"""
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w


def write_table(ws, start_row, headers, rows, col_widths, root_cause_row_indices=None, dropdowns=None):
    """
    在指定起始行写入一个表格（含表头 + 数据行）。
    root_cause_row_indices: 数据行中需要黄色高亮的索引列表（0-based 数据行索引）
    dropdowns: 可选，字典 {列索引(1-based): [选项1, 选项2, ...]}，给指定列加 Excel 数据验证下拉框
    返回：下一可用行号
    """
    if root_cause_row_indices is None:
        root_cause_row_indices = []
    if dropdowns is None:
        dropdowns = {}

    # 表头
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=start_row, column=col_idx, value=header)
        apply_header_style(cell)
    ws.row_dimensions[start_row].height = 28

    # 数据行
    for i, row_data in enumerate(rows):
        current_row = start_row + 1 + i
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=current_row, column=col_idx, value=value)
            is_rc = i in root_cause_row_indices
            apply_body_style(cell, i, is_root_cause=is_rc)
        # 自动行高：根据每个单元格内容和列宽精确计算换行数，确保文字完全可见
        row_height = _calc_row_height(row_data, col_widths)
        ws.row_dimensions[current_row].height = row_height

    # 设置列宽
    set_column_widths(ws, col_widths)

    # 添加数据验证下拉框（如果指定了 dropdowns）
    for col_idx, options in dropdowns.items():
        if col_idx < 1 or col_idx > len(headers):
            continue
        # 构造 Excel 数据验证公式（用英文逗号分隔，最多 255 字符）
        # 如果选项总长超 255，用引用范围方式（这里简化处理，限制选项数）
        options_str = ",".join(options)
        if len(options_str) > 255:
            # 选项太多，跳过（避免 Excel 报错）
            continue
        dv = DataValidation(
            type="list",
            formula1=f'"{options_str}"',
            allow_blank=True,
            showDropDown=False,  # False = 显示下拉箭头（openpyxl 的 showDropDown 是反逻辑）
            showErrorMessage=True,
            errorTitle="输入无效",
            error=f"请从下拉框选择：{', '.join(options)}",
            promptTitle="请选择",
            prompt=f"可选值：{', '.join(options)}",
        )
        # 应用到数据行（不含表头）
        col_letter = get_column_letter(col_idx)
        first_row = start_row + 1
        last_row = start_row + len(rows)
        cell_range = f"{col_letter}{first_row}:{col_letter}{last_row}"
        dv.add(cell_range)
        ws.add_data_validation(dv)

    # 注意：冻结窗格已移到 generate_excel 末尾统一设置（避免多次设置互相覆盖）
    return start_row + 1 + len(rows)


def write_section_title(ws, row, title, span_cols=6):
    """写入章节标题（深蓝底白字，合并单元格）。"""
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span_cols)
    cell = ws.cell(row=row, column=1, value=title)
    apply_header_style(cell)
    ws.row_dimensions[row].height = 30
    return row + 1


def write_yellow_section_title(ws, row, title, span_cols=6):
    """写入黄色章节标题（淡黄底深蓝字，用于区分章节目录层级）。"""
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span_cols)
    cell = ws.cell(row=row, column=1, value=title)
    apply_yellow_section_style(cell)
    ws.row_dimensions[row].height = 30
    return row + 1


def write_kv_block(ws, start_row, kv_pairs, label_col_width=22, value_col_width=40):
    """
    写入键值对块（左标签 + 右值），2 列布局。
    kv_pairs: [(label, value), ...]
    """
    for i, (label, value) in enumerate(kv_pairs):
        current_row = start_row + i
        # 标签列
        label_cell = ws.cell(row=current_row, column=1, value=label)
        label_cell.font = Font(name=FONT_NAME, size=10, bold=True, color="FFFFFF")
        label_cell.fill = get_subheader_fill()
        label_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        label_cell.border = get_thin_border()
        # 值列
        value_cell = ws.cell(row=current_row, column=2, value=value)
        value_cell.font = get_body_font()
        value_cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        value_cell.border = get_thin_border()
        ws.row_dimensions[current_row].height = _calc_row_height(
            [label, value], [label_col_width, value_col_width]
        )

    set_column_widths(ws, [label_col_width, value_col_width])
    return start_row + len(kv_pairs)


def _apply_auto_fill(ws, report_number):
    """自动填充模式：扫描所有单元格，把 ____ 替换为合理示例值。
    
    填充规则：
    - 姓名字段（D1 团队、D8 签名栏）：按角色分配化名
    - 日期字段：基于当前日期计算
    - 联系方式：内部分机号
    - 责任人：按措施类型分配角色
    """
    import datetime
    from openpyxl.utils import get_column_letter
    
    today = datetime.date.today()
    today_str = today.strftime("%Y-%m-%d")
    date_plus = lambda days: (today + datetime.timedelta(days=days)).strftime("%Y-%m-%d")
    
    # 角色对应的化名（固定，方便用户区分）
    ROLE_NAMES = {
        "团队领导（质量工程师）": "张伟",
        "工艺工程师": "李娜",
        "设备工程师": "刘强",
        "生产主管": "陈静",
        "设计工程师": "赵磊",
        "SQE（供应商质量工程师）": "周敏",
        "质量经理（审核）": "王芳",
        "编制（质量工程师）": "张伟",
        "审核（质量经理）": "王芳",
        "批准（质量总监）": "孙健",
        "客户确认（如需）": "____",
    }
    
    # 部门对应的化名（用于 D3/D5/D6/D7 责任人列）
    DEPT_PERSON = {
        "质量部": "张伟",
        "工艺部": "李娜",
        "设备部": "刘强",
        "生产部": "陈静",
        "研发部": "赵磊",
        "物流部": "周敏",
        "销售部": "吴洋",
    }
    
    # 扫描所有单元格，按上下文判断 ____ 应该填什么
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            if cell.value is None:
                continue
            val = str(cell.value).strip()
            if val != "____" and not val.startswith("____（") and not val.startswith("请填写"):
                continue
            
            # 获取同行左侧单元格内容，用于上下文判断
            left_cell = ws.cell(row=cell.row, column=cell.col_idx - 1) if cell.col_idx > 1 else None
            left_val = str(left_cell.value) if left_cell and left_cell.value else ""
            
            # 获取本行第一列（通常是序号或角色）
            first_col_val = str(ws.cell(row=cell.row, column=1).value or "")
            
            # 判断字段类型并填充
            # 先尝试通用填充
            replacement = _guess_fill_value(
                left_val=left_val,
                first_col_val=first_col_val,
                col_idx=cell.col_idx,
                role_names=ROLE_NAMES,
                dept_person=DEPT_PERSON,
                today_str=today_str,
                date_plus=date_plus,
                report_number=report_number,
            )
            
            # 如果通用填充没匹配，尝试根据 val 内容激进填充
            if not replacement and val:
                replacement = _aggressive_fill(val, left_val, first_col_val, today_str, date_plus)
            
            if replacement:
                cell.value = replacement


def _apply_auto_fill_word(doc, report_number):
    """自动填充模式（Word 版）：遍历所有表格的所有单元格，把 ____ 替换为合理示例值。
    
    复用 _guess_fill_value 的判断逻辑，但针对 Word 表格结构调整：
    - Word 表格每行有多个 cell，第一个 cell 通常是标签或角色名
    - 通过 cell 所在列索引和同行第一个 cell 的内容判断该填什么
    """
    import datetime
    
    today = datetime.date.today()
    today_str = today.strftime("%Y-%m-%d")
    date_plus = lambda days: (today + datetime.timedelta(days=days)).strftime("%Y-%m-%d")
    
    # 角色对应的化名（与 Excel 版保持一致）
    ROLE_NAMES = {
        "团队领导（质量工程师）": "张伟",
        "工艺工程师": "李娜",
        "设备工程师": "刘强",
        "生产主管": "陈静",
        "设计工程师": "赵磊",
        "SQE（供应商质量工程师）": "周敏",
        "质量经理（审核）": "王芳",
        "编制（质量工程师）": "张伟",
        "审核（质量经理）": "王芳",
        "批准（质量总监）": "孙健",
        "客户确认（如需）": "____",
    }
    
    DEPT_PERSON = {
        "质量部": "张伟",
        "工艺部": "李娜",
        "设备部": "刘强",
        "生产部": "陈静",
        "研发部": "赵磊",
        "物流部": "周敏",
        "销售部": "吴洋",
    }
    
    # 联系方式分机号映射
    EXT_MAP = {"张伟": "8001", "李娜": "8002", "刘强": "8003", "陈静": "8004",
               "赵磊": "8005", "周敏": "8006", "王芳": "8007", "孙健": "8008", "吴洋": "8009"}
    
    # 责任人轮换
    PERSON_ROTATION = ["张伟", "李娜", "刘强", "陈静", "周敏", "张伟", "王芳"]
    
    # 遍历文档中所有表格
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if not cells:
                continue
            
            # 获取本行第一个 cell 的内容（用于上下文判断）
            first_cell_text = cells[0].text.strip() if cells[0].text else ""
            
            for col_idx, cell in enumerate(cells, start=1):
                cell_text = cell.text.strip() if cell.text else ""
                
                # 只处理 ____ 或以 ____（ 开头的单元格
                if cell_text != "____" and not cell_text.startswith("____（") and not cell_text.startswith("请填写"):
                    continue
                
                # 获取左侧 cell 的内容
                left_val = cells[col_idx - 2].text.strip() if col_idx > 1 and col_idx - 2 < len(cells) else ""
                
                replacement = _guess_fill_value(
                    left_val=left_val,
                    first_col_val=first_cell_text,
                    col_idx=col_idx,
                    role_names=ROLE_NAMES,
                    dept_person=DEPT_PERSON,
                    today_str=today_str,
                    date_plus=date_plus,
                    report_number=report_number,
                )
                
                # 如果通用填充没匹配，尝试根据 cell_text 内容激进填充
                if not replacement and cell_text:
                    replacement = _aggressive_fill(cell_text, left_val, first_cell_text, today_str, date_plus)
                
                if replacement:
                    # 清空 cell 并写入新值（用更健壮的方式）
                    # 方法：清空所有 paragraph 的所有 run，然后在第一个 paragraph 写入新 run
                    try:
                        # 清空第一个 paragraph 的所有 run
                        if cell.paragraphs:
                            first_para = cell.paragraphs[0]
                            # 删除所有现有 run
                            for run in first_para.runs:
                                run.text = ""
                            # 如果没有 run，创建一个
                            if not first_para.runs:
                                first_para.add_run(replacement)
                            else:
                                # 在第一个 run 写入新值
                                first_para.runs[0].text = replacement
                            # 删除其他 paragraph
                            for p in cell.paragraphs[1:]:
                                p_element = p._element
                                p_element.getparent().remove(p_element)
                        else:
                            # 没有 paragraph，直接用 cell.text
                            cell.text = replacement
                    except Exception as e:
                        # 兜底：用 cell.text
                        try:
                            cell.text = replacement
                        except Exception:
                            pass


def _aggressive_fill(val, left_val, first_col_val, today_str, date_plus):
    """激进填充模式：对 _guess_fill_value 无法处理的占位符，根据内容直接填示例值。
    用户启用 auto_fill 时调用，确保所有空白都被填上。
    """
    val_clean = val.strip()
    left_clean = left_val.strip() if left_val else ""
    first_clean = first_col_val.strip() if first_col_val else ""
    
    # D4 5Why 答案（"请填写：..." 开头）
    if val_clean.startswith("请填写：方向"):
        return "方向①设备/工装异常：物流转运过程外力冲击导致凹陷（示例）"
    if val_clean.startswith("请填写：基于 Why 1"):
        return "包装方案未针对轮毂形状设计有效缓冲结构（示例）"
    if val_clean.startswith("请填写：定位到具体的监控点"):
        return "包装验证职责归属不明确，工艺/物流/采购互相推诿，监控点失效（示例）"
    if val_clean.startswith("请填写：定位到具体的 SOP"):
        return "部门职责划分文件中未明确包装验证责任归属，SOP存在空白（示例）"
    if val_clean.startswith("请填写：定位到体系/规范"):
        return "管理制度不透明，未建立R&R职责矩阵，跨部门协作事项无人负责（示例）"
    
    # D4 6M 判定列
    if val_clean.startswith("____（根因/贡献因子"):
        if "法" in left_clean:
            return "根因"
        if "人" in left_clean:
            return "贡献因子"
        return "排除"
    
    # D4 RC1/RC2/RC3 总结
    if val_clean.startswith("请填写直接原因"):
        return "包装缓冲不足，物流转运外力冲击导致凹陷（示例）"
    if val_clean.startswith("请填写管理原因"):
        return "包装验证职责不清，跨部门协作机制缺失（示例）"
    if val_clean.startswith("请填写系统原因"):
        return "管理制度不透明，未建立R&R职责矩阵（示例）"
    
    # Where 字段
    if val_clean.startswith("____（工序"):
        return "客户IQC来料检验（示例）"
    
    # 兜底：所有剩余 ____ 都填示例标记
    if val_clean == "____":
        return "（示例数据，请替换为实际值）"
    if val_clean.startswith("____（"):
        import re as _re
        m = _re.search(r'____（([^）]+)', val_clean)
        if m:
            hint = m.group(1)
            return f"（示例：{hint}）"
        return "（示例数据）"
    
    return None


def _guess_fill_value(left_val, first_col_val, col_idx, role_names, dept_person, today_str, date_plus, report_number):
    """根据上下文猜测 ____ 应该填什么值。返回 None 表示无法判断，保留 ____"""
    left_val_clean = left_val.strip()
    first_col_clean = first_col_val.strip()
    
    # 1. 姓名/签名栏：用本行第一列判断是否是角色名行（D1 团队表 / D8 签名栏）
    # D1 团队表：第1列=角色名，第2列=姓名，第3列=部门，第4列=联系方式
    # D8 签名栏：第1列=角色名，第2列=姓名，第3列=签名，第4列=日期
    if first_col_clean in role_names:
        name = role_names[first_col_clean]
        if name != "____":
            # 判断是 D1 团队表还是 D8 签名栏
            # D8 角色名包含"编制"/"审核"/"批准"/"客户确认"
            # D8 签名栏的角色名以"编制"/"审核"/"批准"/"客户确认"开头
            # D1 团队表的角色名"质量经理（审核）"包含"审核"但不在开头
            is_d8_signature = any(first_col_clean.startswith(kw) for kw in ["编制", "审核", "批准", "客户确认"])
            
            if col_idx == 2:
                return name  # 姓名
            if is_d8_signature:
                # D8 签名栏
                if col_idx == 3:
                    return name  # 签名 = 姓名
                if col_idx == 4:
                    return today_str  # 日期 = 当天
            else:
                # D1 团队表
                if col_idx == 4:
                    # 联系方式：内部分机号
                    ext_map = {"张伟": "8001", "李娜": "8002", "刘强": "8003", "陈静": "8004",
                              "赵磊": "8005", "周敏": "8006", "王芳": "8007", "孙健": "8008", "吴洋": "8009"}
                    return ext_map.get(name, "____")
    
    # 2. D3/D5/D6/D7 责任人列：第一列是序号(1/2/3...)，按行号分配角色
    if first_col_clean in ["1", "2", "3", "4", "5", "6", "7"]:
        seq = int(first_col_clean)
        # 责任人轮换分配（质量/工艺/设备/生产/物流/质量/质量）
        person_rotation = ["张伟", "李娜", "刘强", "陈静", "周敏", "张伟", "王芳"]
        if col_idx == 3:  # D3 责任人列（序号/措施/责任人/完成时间/验证/状态）
            return person_rotation[min(seq - 1, len(person_rotation) - 1)]
        if col_idx == 4:  # D3 完成时间列
            return date_plus(2 + seq)  # 2/3/4/5/6 天后
        if col_idx == 5:  # D3 验证方法列
            return "100%全检记录：检验XX件，发现不良XX件（示例）"
        # D5 表格列: 序号/CA方案/针对根因/可行性/风险评估/决策
        # D6 表格列: 序号/实施措施/目标根因/责任人/完成时间/状态/验证结果
        # D7 表格列: 序号/横向展开措施/推广范围/责任人/完成时间/状态
        # 判断当前在哪个表，用 col_idx 推断
    
    # 3. D6 责任人列（第4列）、完成时间列（第5列）
    if first_col_clean in ["1", "2", "3", "4", "5", "6", "7"]:
        seq = int(first_col_clean)
        person_rotation = ["张伟", "李娜", "刘强", "陈静", "周敏", "张伟", "王芳"]
        if col_idx == 4:  # D6 责任人列
            return person_rotation[min(seq - 1, len(person_rotation) - 1)]
        if col_idx == 5:  # D6 完成时间列
            return date_plus(7 + seq * 3)  # 7/10/13/16/19 天后
        if col_idx == 7:  # D6 验证结果列
            return "已完成验证，连续30天数据达标（示例）"
        # D7 责任人列（第4列）、完成时间列（第5列）
        if col_idx == 4:
            return person_rotation[min(seq - 1, len(person_rotation) - 1)]
        if col_idx == 5:
            return date_plus(14 + seq * 7)  # 14/21/28/35 天后
    
    # 4. D0-D2 信息表的空白字段
    if left_val_clean == "报告发起人":
        return "张伟"
    if left_val_clean == "客户联系人":
        return "李工（IQC主管）"
    if left_val_clean == "客户投诉日期":
        return date_plus(-3)  # 3 天前
    if left_val_clean == "客户投诉单号":
        return f"CC-{today_str.replace('-', '')}-{1000 + hash(first_col_val) % 9000}"
    if left_val_clean == "客户反馈渠道":
        return "邮件"
    if left_val_clean == "产品编号 / 零件号":
        return "WH-2026-001（示例）"
    if left_val_clean == "批次号":
        return f"B{today_str.replace('-', '')}01"
    if left_val_clean == "不良数量":
        return "12 件（示例）"
    if left_val_clean == "8D 负责人":
        return "张伟"
    
    # 5. D2 5W2H 表的空白
    if left_val_clean == "When（何时发现）":
        return f"{date_plus(-3)} 08:30 白班"
    if left_val_clean == "Who（谁发现的）":
        return "客户 IQC 检验员"
    if left_val_clean == "Why（为什么是问题）":
        return "违反客户外观标准，影响产品功能（示例）"
    if left_val_clean == "How（如何发现）":
        return "100% 外观目视检查"
    
    # 6. D3 遏制有效性验证表的空白
    if left_val_clean == "遏制开始日期":
        return date_plus(-2)
    if left_val_clean == "遏制后不良率（24h内）":
        return "0 PPM（示例，已隔离不良品）"
    if left_val_clean == "遏制后不良率（72h内）":
        return "0 PPM（示例，已隔离不良品）"
    if left_val_clean == "遏制结论":
        return "达标，遏制后不良率降至0 PPM（示例）"
    if left_val_clean == "遏制措施截止日期":
        return date_plus(7)
    if left_val_clean == "客户是否认可":
        return "待确认"
    
    # 7. D4 根本原因验证的空白
    if left_val_clean == "验证方法":
        return "现场观察 + 数据收集"
    if left_val_clean == "验证数据":
        return "连续30天数据，不良率从500 PPM降至50 PPM（示例）"
    if left_val_clean == "验证结论":
        return "待进一步验证"
    if left_val_clean == "验证人 / 日期":
        return f"张伟 / {today_str}"
    
    # 8. D7 PFMEA 更新表的空白
    if left_val_clean == "本次失效模式":
        return "轮毂凹陷（包装缓冲不足导致）（示例）"
    if left_val_clean == "原 RPN":
        return "270（S=9, O=5, D=6）（示例）"
    if left_val_clean == "更新后 RPN":
        return "80（S=9, O=2, D=4）（示例）"
    if left_val_clean == "PFMEA 更新人 / 日期":
        return f"李娜 / {today_str}"
    
    # 9. D8 关闭确认的空白
    if left_val_clean == "所有 CA 是否实施完成":
        return "否（进行中）"
    if left_val_clean == "所有 CA 是否验证有效":
        return "否（待验证）"
    if left_val_clean == "客户是否认可":
        return "待确认"
    if left_val_clean == "所有文件是否更新（SOP/PFMEA/培训教材）":
        return "否（待更新）"
    if left_val_clean == "横向展开是否完成":
        return "否（待展开）"
    if left_val_clean == "关闭日期":
        return "2026-07-30（示例）"
    if left_val_clean == "关闭结论":
        return "暂不关闭，待所有 CA 实施并验证"
    if left_val_clean == "本次问题处理经验":
        return "包装验证职责需在R&R矩阵中明确，避免跨部门推诿（示例）"
    if left_val_clean == "可改进之处":
        return "包装验证职责需在R&R矩阵中明确，避免跨部门推诿（示例）"
    if left_val_clean == "对其他产品的启示":
        return "包装验证职责需在R&R矩阵中明确，避免跨部门推诿（示例）"
    if left_val_clean == "建议改进的管理流程":
        return "包装验证职责需在R&R矩阵中明确，避免跨部门推诿（示例）"
    
    # 10. D5/D6 验证结论的空白
    if left_val_clean == "有效性验证":
        return "需至少30天连续数据，不良率降至目标值50 PPM以下（示例）"
    
    # 无法判断的字段，保留 ____
    return None


# ============================================================
# Excel 生成（单 Sheet 版本：D0-D8 合并为一个表格）
# ============================================================

def generate_excel(context, template, output_path, report_number, auto_fill=False):
    """生成 Excel 文件（单 Sheet：8D报告）。"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "8D报告"
    ws.sheet_properties.tabColor = "003366"

    row = 1
    d0_d2 = template.get("d0_d2", {})

    # ==================== D0-D2 ====================
    row = write_yellow_section_title(ws, row, "D0-D2  问题基本信息与团队组建", span_cols=4)
    row += 1

    # 报告元信息
    row = write_kv_block(ws, row, [
        ("8D 报告编号", report_number),
        ("报告生成日期", datetime.datetime.now().strftime("%Y-%m-%d")),
        ("报告发起人", "____"),
        ("报告状态", "进行中"),
    ])
    row += 1

    # 客户信息
    row = write_section_title(ws, row, "一、客户信息", span_cols=2)
    row = write_kv_block(ws, row, [
        ("客户名称", context.get("customer", "____")),
        ("客户联系人", "____"),
        ("客户投诉日期", "____"),
        ("客户投诉单号", "____"),
        ("客户反馈渠道", "____（邮件 / 电话 / 8D通知单 / 现场）"),
    ])
    row += 1

    # 产品信息
    row = write_section_title(ws, row, "二、产品信息", span_cols=2)
    row = write_kv_block(ws, row, [
        ("产品名称", context.get("product", "____")),
        ("产品编号 / 零件号", "____"),
        ("批次号", "____"),
        ("批量数量", context.get("batch_size", "____")),
        ("不良数量", "____"),
        ("不良率", context.get("defect_rate", "____")),
        ("缺陷等级", d0_d2.get("defect_level_hint", "____")),
        ("发现位置", d0_d2.get("discovery_location_hint", "____")),
        ("影响", d0_d2.get("impact_hint", "____")),
    ])
    row += 1

    # D2 5W2H
    row = write_section_title(ws, row, "三、D2 问题描述（5W2H）", span_cols=2)
    row = write_kv_block(ws, row, [
        ("What（什么问题）", context.get("defect", "____")),
        ("When（何时发现）", "____（日期 / 班次）"),
        ("Where（何处发现）", "____（工序 / 客户环节 / 产品位置）"),
        ("Who（谁发现的）", "____"),
        ("Why（为什么是问题）", "____（违反的标准 / 规格要求）"),
        ("How（如何发现）", "____（检验方法 / 测试方法）"),
        ("How many（不良数量/率）", f"{context.get('defect_rate', '____')}，批量 {context.get('batch_size', '____')}"),
    ])
    row += 1

    # 问题陈述
    row = write_section_title(ws, row, "四、问题陈述", span_cols=2)
    problem_stmt = (
        f"产品「{context.get('product', '____')}」在 {d0_d2.get('discovery_location_hint', '____')} "
        f"发现「{context.get('defect', '____')}」问题，"
        f"不良率 {context.get('defect_rate', '____')}，"
        f"涉及批次数量 {context.get('batch_size', '____')}。"
        f"该问题已导致 {d0_d2.get('impact_hint', '____')}，需要立即启动 8D 问题解决流程。"
    )
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
    cell = ws.cell(row=row, column=1, value=problem_stmt)
    cell.font = get_body_font()
    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    cell.border = get_thin_border()
    ws.row_dimensions[row].height = _calc_row_height([problem_stmt], [62])  # 合并列总宽 ≈ 22+40
    row += 1

    # D1 团队
    row += 1
    row = write_section_title(ws, row, "五、D1 团队组建表", span_cols=4)
    row = write_table(ws, row,
        ["角色", "姓名", "部门", "联系方式"],
        [
            ["团队领导（质量工程师）", "____", "质量部", "____"],
            ["工艺工程师", "____", "工艺部", "____"],
            ["设备工程师", "____", "设备部", "____"],
            ["生产主管", "____", "生产部", "____"],
            ["设计工程师", "____", "研发部", "____"],
            ["SQE（供应商质量工程师）", "____", "质量部", "____"],
            ["质量经理（审核）", "____", "质量部", "____"],
        ],
        col_widths=[28, 18, 14, 22],
    )

    # ==================== D3 ====================
    row += 2
    row = write_yellow_section_title(ws, row, "D3  临时遏制措施（Interim Containment Actions）", span_cols=6)
    row += 1
    row = write_section_title(ws, row, "一、遏制措施清单", span_cols=6)
    d3 = template.get("d3_template", {})
    actions = d3.get("containment_actions", [])
    while len(actions) < 5:
        actions.append("____（请补充遏制措施）")
    row = write_table(ws, row,
        ["序号", "遏制措施", "责任人", "完成时间", "验证方法", "状态"],
        [[i + 1, actions[i], "____", "____", "____（如100%全检记录/不良率对比）", "待执行"] for i in range(5)],
        col_widths=[6, 50, 12, 14, 28, 10],
        dropdowns={6: ["待执行", "进行中", "已完成", "延期", "取消"]},
    )
    row += 1
    row = write_section_title(ws, row, "二、遏制有效性验证", span_cols=2)
    row = write_kv_block(ws, row, [
        ("遏制前不良率", context.get("defect_rate", "____")),
        ("遏制开始日期", "____"),
        ("遏制后不良率（24h内）", "____"),
        ("遏制后不良率（72h内）", "____"),
        ("遏制结论", "____（达标 / 未达标，是否需要继续遏制）"),
        ("遏制措施截止日期", "____"),
        ("客户是否认可", "____"),
    ])

    # ==================== D4 ====================
    row += 2
    row = write_yellow_section_title(ws, row, "D4  根本原因分析（Root Cause Analysis）", span_cols=4)
    row += 1
    row = write_section_title(ws, row, "一、5Why 分析", span_cols=4)
    d4 = template.get("d4_template", {})
    five_why = d4.get("5why_path", {})
    steps = five_why.get("steps", [])
    why_rows = []
    root_cause_indices = []

    if steps:
        for idx, step in enumerate(steps):
            level = step.get("level", "____")
            question = step.get("question", "____")
            answer = step.get("answer", "____")
            evidence = step.get("evidence", "____")
            why_rows.append([level, question, answer, evidence])
            if "根因" in str(level):
                root_cause_indices.append(idx)
    else:
        problem = five_why.get("problem", "____")
        why_rows.append(["问题", problem, "—", "—"])
        why1 = five_why.get("why1", "____")
        why_rows.append(["Why 1", "为什么出现该问题？", why1, "____"])
        why2_hints = five_why.get("why2_hints", [])
        why_rows.append(["Why 2", f"为什么：{why1[:30]}...", "\n".join(f"• {h}" for h in why2_hints) if why2_hints else "____", "____"])
        why3_hints = five_why.get("why3_hints", [])
        why_rows.append(["Why 3", "为什么会出现上述原因？", "\n".join(f"• {h}" for h in why3_hints) if why3_hints else "____", "____"])
        why4_hints = five_why.get("why4_hints", [])
        why_rows.append(["Why 4", "为什么管理/流程未能预防？", "\n".join(f"• {h}" for h in why4_hints) if why4_hints else "____", "____"])
        why5_root = five_why.get("why5_root", "____")
        why_rows.append(["Why 5（根因）", "为什么流程/规范存在缺陷？", why5_root, "____"])
        root_cause_indices = [5]

    row = write_table(ws, row,
        ["层级", "问题", "答案", "证据"],
        why_rows,
        col_widths=[14, 38, 50, 22],
        root_cause_row_indices=root_cause_indices,
    )

    row += 1
    row = write_section_title(ws, row, "二、鱼骨图 6M 排查", span_cols=4)
    six_m = d4.get("6m_analysis", {})

    if isinstance(six_m, list):
        m_rows = []
        m_rc_indices = []
        for idx, item in enumerate(six_m):
            m_name = item.get("m", "____")
            finding = item.get("finding", "____")
            judgment = item.get("judgment", "____")
            m_rows.append([m_name, finding, judgment])
            if "根本原因" in str(judgment):
                m_rc_indices.append(idx)
        row = write_table(ws, row,
            ["6M 维度", "排查结果", "判定"],
            m_rows,
            col_widths=[18, 55, 18],
            root_cause_row_indices=m_rc_indices,
            dropdowns={3: ["根因", "贡献因子", "排除", "待确认"]},
        )
    else:
        row = write_table(ws, row,
            ["6M 维度", "候选原因", "证据", "是否根因"],
            [
                ["Man（人）", six_m.get("man", "____"), "____", "否"],
                ["Machine（机）", six_m.get("machine", "____"), "____", "否"],
                ["Material（料）", six_m.get("material", "____"), "____", "否"],
                ["Method（法）", six_m.get("method", "____"), "____", "否"],
                ["Measurement（测）", six_m.get("measurement", "____"), "____", "否"],
                ["Environment（环）", six_m.get("environment", "____"), "____", "否"],
            ],
            col_widths=[16, 50, 24, 14],
            dropdowns={4: ["是", "否", "待确认"]},
        )

    row += 1
    row = write_section_title(ws, row, "三、根本原因总结", span_cols=2)
    rc_summary = d4.get("root_cause_summary", [])
    if rc_summary:
        rc_kv = [(f"{rc.get('id', '____')}（{rc.get('type', '____')}）", rc.get("description", "____")) for rc in rc_summary]
        row = write_kv_block(ws, row, rc_kv, label_col_width=22, value_col_width=70)
    else:
        row = write_kv_block(ws, row, [
            ("RC1（直接原因）", "____（参考 5Why 第 1-2 层结论）"),
            ("RC2（管理原因）", "____（参考 5Why 第 3-4 层结论）"),
            ("RC3（系统原因）", "____（参考 5Why 第 5 层结论）"),
        ], label_col_width=22, value_col_width=70)

    row += 1
    row = write_section_title(ws, row, "四、根本原因验证结论", span_cols=2)
    verify_text = d4.get("verification", "")
    if verify_text:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
        cell = ws.cell(row=row, column=1, value=verify_text)
        cell.font = get_body_font()
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        cell.border = get_thin_border()
        ws.row_dimensions[row].height = _calc_row_height([verify_text], [92])  # 合并列总宽 ≈ 22+70
    else:
        row = write_kv_block(ws, row, [
            ("验证方法", "____（现场观察 / 数据收集 / 重现试验）"),
            ("验证数据", "____"),
            ("验证结论", "____（已确认 / 待进一步验证）"),
            ("验证人 / 日期", "____"),
        ], label_col_width=22, value_col_width=70)

    # ==================== D5-D6 ====================
    row += 2
    row = write_yellow_section_title(ws, row, "D5-D6  永久纠正措施（Permanent Corrective Actions）", span_cols=7)
    row += 1
    row = write_section_title(ws, row, "一、D5 CA 方案评估矩阵", span_cols=6)
    d5_d6 = template.get("d5_d6_template", {})
    permanent_actions = d5_d6.get("permanent_actions", [])

    ca_rows = []
    for i, action in enumerate(permanent_actions, start=1):
        if isinstance(action, dict):
            ca_rows.append([i, action.get("action", "____"), action.get("target", "____"), "____（高/中/低）", "____（是否引入新风险）", "____（采纳/否决）"])
        else:
            ca_rows.append([i, str(action), "____", "____（高/中/低）", "____（是否引入新风险）", "____（采纳/否决）"])
    while len(ca_rows) < 3:
        ca_rows.append([len(ca_rows) + 1, "____（请补充 CA 方案）", "____", "____", "____", "____"])

    row = write_table(ws, row,
        ["序号", "CA 方案", "针对根因", "可行性", "风险评估", "决策"],
        ca_rows,
        col_widths=[6, 50, 14, 12, 30, 12],
        dropdowns={4: ["高", "中", "低"], 6: ["采纳", "否决", "待评估"]},
    )

    row += 1
    row = write_section_title(ws, row, "二、D6 实施计划跟踪表", span_cols=7)
    d6_rows = []
    for i, action in enumerate(permanent_actions, start=1):
        if isinstance(action, dict):
            d6_rows.append([i, action.get("action", "____"), action.get("target", "____"), action.get("responsible", "____"), action.get("due_date", "____"), "____（未开始/进行中/已完成）", "____"])
        else:
            d6_rows.append([i, str(action), "____", "____", "____", "____（未开始/进行中/已完成）", "____"])
    while len(d6_rows) < 3:
        d6_rows.append([len(d6_rows) + 1, "____", "____", "____", "____", "____", "____"])

    row = write_table(ws, row,
        ["序号", "实施措施", "目标根因", "责任人", "完成时间", "状态", "验证结果"],
        d6_rows,
        col_widths=[6, 40, 12, 14, 14, 18, 22],
        dropdowns={6: ["未开始", "进行中", "已完成", "延期", "取消"]},
    )

    # ==================== D7 ====================
    row += 2
    row = write_yellow_section_title(ws, row, "D7  预防再发生 — 横向展开（Prevent Recurrence — Yokoten）", span_cols=6)
    row += 1
    row = write_section_title(ws, row, "一、横向展开措施清单", span_cols=5)
    d7 = template.get("d7_template", {})
    yokoten = d7.get("yokoten", [])
    y_rows = [[i + 1, item, "____（同类产线/产品/客户）", "____", "____", "____"] for i, item in enumerate(yokoten)]
    while len(y_rows) < 4:
        y_rows.append([len(y_rows) + 1, "____（请补充 Yokoten 措施）", "____", "____", "____", "____"])
    row = write_table(ws, row,
        ["序号", "横向展开措施", "推广范围", "责任人", "完成时间", "状态"],
        y_rows,
        col_widths=[6, 50, 22, 14, 14, 14],
        dropdowns={6: ["未开始", "进行中", "已完成", "延期", "不适用"]},
    )

    row += 1
    row = write_section_title(ws, row, "二、PFMEA 更新", span_cols=2)
    row = write_kv_block(ws, row, [
        ("本次失效模式", "____"),
        ("原 RPN", "____"),
        ("更新后 RPN", "____"),
        ("PFMEA 更新人 / 日期", "____"),
    ], label_col_width=24, value_col_width=60)

    # ==================== D8 ====================
    row += 2
    row = write_yellow_section_title(ws, row, "D8  团队认可与关闭（Team Recognition & Closure）", span_cols=4)
    row += 1
    row = write_section_title(ws, row, "一、关闭确认", span_cols=2)
    row = write_kv_block(ws, row, [
        ("所有 CA 是否实施完成", "____（是/否）"),
        ("所有 CA 是否验证有效", "____（是/否）"),
        ("客户是否认可", "____（是/否，附客户确认邮件）"),
        ("所有文件是否更新（SOP/PFMEA/培训教材）", "____（是/否）"),
        ("横向展开是否完成", "____（是/否）"),
        ("关闭日期", "____"),
        ("关闭结论", "____（同意关闭 / 暂不关闭，原因：____）"),
    ], label_col_width=34, value_col_width=60)

    row += 1
    row = write_section_title(ws, row, "二、经验教训", span_cols=2)
    row = write_kv_block(ws, row, [
        ("本次问题处理经验", "____"),
        ("可改进之处", "____"),
        ("对其他产品的启示", "____"),
        ("建议改进的管理流程", "____"),
    ], label_col_width=24, value_col_width=60)

    row += 1
    row = write_section_title(ws, row, "三、签名栏", span_cols=4)
    row = write_table(ws, row,
        ["角色", "姓名", "签名", "日期"],
        [
            ["编制（质量工程师）", "____", "____", "____"],
            ["审核（质量经理）", "____", "____", "____"],
            ["批准（质量总监）", "____", "____", "____"],
            ["客户确认（如需）", "____", "____", "____"],
        ],
        col_widths=[24, 20, 25, 18],
    )

    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
    cell = ws.cell(row=row, column=1, value=f"8D 报告编号：{report_number}")
    cell.font = Font(name=FONT_NAME, size=10, italic=True, color="666666")
    cell.alignment = Alignment(horizontal="right", vertical="center")

    # 冻结窗格：只冻结第 1 行（标题行），不冻结列
    # 这样用户向下滚动时标题始终可见，且不会锁定中间表格的滚动
    ws.freeze_panes = "A2"

    # 自动填充模式：把所有 ____ 替换为合理示例值
    if auto_fill:
        _apply_auto_fill(ws, report_number)
        print(f"[INFO] 已启用自动填充模式，所有 ____ 已替换为示例值")

    # ── 最终：重新计算所有行的行高（确保 auto_fill 后行高也正确）──
    _recalc_all_row_heights(ws)

    wb.save(output_path)
    print(f"[OK] Excel 已生成：{output_path}")


# ============================================================
# Word 生成
# ============================================================

def set_cell_bg(cell, color_hex):
    """设置 Word 表格单元格背景色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    """设置 Word 表格单元格四边框。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_run_font(run, font_name=FONT_NAME, size=10.5, bold=False, color=None):
    """设置 run 的字体（中文 + 英文）。"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # 中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_paragraph(doc, text, bold=False, size=10.5, color=None, alignment=None, indent=False):
    """添加段落。"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    """添加章节标题（自定义样式，避免依赖 Word 默认 Heading 样式字体问题）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=14, bold=True, color=HEADER_FILL)
    elif level == 2:
        set_run_font(run, size=12, bold=True, color=SUBHEADER_FILL)
    else:
        set_run_font(run, size=11, bold=True, color="333333")
    return p


def add_table(doc, headers, rows, col_widths_cm=None, root_cause_row_indices=None):
    """
    添加带格式的表格。
    headers: [str]
    rows: [[str]]
    col_widths_cm: [float]
    root_cause_row_indices: list of 0-based row indices that should be yellow-highlighted
    """
    if root_cause_row_indices is None:
        root_cause_row_indices = []

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 设置列宽
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        set_cell_bg(cell, HEADER_FILL)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=10.5, bold=True, color="FFFFFF")

    # 数据行
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        is_rc = r_idx in root_cause_row_indices
        for c_idx, value in enumerate(row_data):
            cell = row_cells[c_idx]
            cell.text = ""
            if is_rc:
                set_cell_bg(cell, ROOT_CAUSE_FILL)
            elif r_idx % 2 == 1:
                set_cell_bg(cell, ALT_ROW_FILL)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_run_font(run, size=10, bold=is_rc)

    return table


def add_kv_table(doc, kv_pairs, label_width_cm=4.5, value_width_cm=12):
    """添加键值对表格（2 列）。"""
    table = doc.add_table(rows=len(kv_pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, (label, value) in enumerate(kv_pairs):
        # 标签列
        label_cell = table.rows[i].cells[0]
        label_cell.width = Cm(label_width_cm)
        label_cell.text = ""
        set_cell_bg(label_cell, SUBHEADER_FILL)
        set_cell_borders(label_cell)
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = label_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        set_run_font(run, size=10, bold=True, color="FFFFFF")

        # 值列
        value_cell = table.rows[i].cells[1]
        value_cell.width = Cm(value_width_cm)
        value_cell.text = ""
        if i % 2 == 1:
            set_cell_bg(value_cell, ALT_ROW_FILL)
        set_cell_borders(value_cell)
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = value_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(value))
        set_run_font(run, size=10)

    return table


def set_doc_default_font(doc):
    """设置文档默认字体为微软雅黑。"""
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)


def set_page_header(doc, report_number):
    """设置页眉，右侧标注 8D-编号。"""
    section = doc.sections[0]
    header = section.header
    # 清空默认段落
    for p in header.paragraphs:
        p.text = ""
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # report_number 已包含 8D- 前缀，无需重复添加
    run = p.add_run(report_number)
    set_run_font(run, size=9, color="666666")


def generate_word(context, template, output_path, report_number, auto_fill=False):
    """生成 Word 文档。"""
    doc = Document()

    # 设置默认字体
    set_doc_default_font(doc)

    # 页面设置：A4
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    # 页眉
    set_page_header(doc, report_number)

    # 主标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run("8D 问题解决报告")
    set_run_font(title_run, size=20, bold=True, color=HEADER_FILL)

    # 报告编号
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(6)
    sub_run = sub_p.add_run(f"报告编号：{report_number}    生成日期：{datetime.datetime.now().strftime('%Y-%m-%d')}")
    set_run_font(sub_run, size=10, color="666666")

    # ============ 一、D0-D2 问题基本信息 ============
    add_heading(doc, "一、D0-D2 问题基本信息与团队组建", level=1)

    d0_d2 = template.get("d0_d2", {})

    add_heading(doc, "1.1 报告基本信息", level=2)
    add_kv_table(doc, [
        ("8D 报告编号", report_number),
        ("报告生成日期", datetime.datetime.now().strftime("%Y-%m-%d")),
        ("报告发起人", "____"),
        ("报告状态", "进行中"),
    ])

    add_heading(doc, "1.2 客户信息", level=2)
    add_kv_table(doc, [
        ("客户名称", context.get("customer", "____")),
        ("客户联系人", "____"),
        ("客户投诉日期", "____"),
        ("客户投诉单号", "____"),
        ("客户反馈渠道", "____（邮件 / 电话 / 8D 通知单 / 现场）"),
    ])

    add_heading(doc, "1.3 产品信息", level=2)
    add_kv_table(doc, [
        ("产品名称", context.get("product", "____")),
        ("产品编号 / 零件号", "____"),
        ("批次号", "____"),
        ("批量数量", context.get("batch_size", "____")),
        ("不良数量", "____"),
        ("不良率", context.get("defect_rate", "____")),
        ("缺陷等级", d0_d2.get("defect_level_hint", "____")),
        ("发现位置", d0_d2.get("discovery_location_hint", "____")),
        ("影响", d0_d2.get("impact_hint", "____")),
    ])

    add_heading(doc, "1.4 D2 问题描述（5W2H）", level=2)
    add_kv_table(doc, [
        ("What（什么问题）", context.get("defect", "____")),
        ("When（何时发现）", "____（日期 / 班次）"),
        ("Where（何处发现）", "____（工序 / 客户环节 / 产品位置）"),
        ("Who（谁发现的）", "____"),
        ("Why（为什么是问题）", "____（违反的标准 / 规格要求）"),
        ("How（如何发现）", "____（检验方法 / 测试方法）"),
        ("How many（不良数量/率）", f"{context.get('defect_rate', '____')}，批量 {context.get('batch_size', '____')}"),
    ])

    add_heading(doc, "1.5 问题陈述", level=2)
    problem_stmt = (
        f"产品「{context.get('product', '____')}」在 {d0_d2.get('discovery_location_hint', '____')} "
        f"发现「{context.get('defect', '____')}」问题，"
        f"不良率 {context.get('defect_rate', '____')}，"
        f"涉及批次数量 {context.get('batch_size', '____')}。"
        f"该问题已导致 {d0_d2.get('impact_hint', '____')}，需要立即启动 8D 问题解决流程。"
    )
    add_paragraph(doc, problem_stmt)

    add_heading(doc, "1.6 D1 团队组建表", level=2)
    add_table(
        doc,
        headers=["角色", "姓名", "部门", "联系方式"],
        rows=[
            ["团队领导（质量工程师）", "____", "质量部", "____"],
            ["工艺工程师", "____", "工艺部", "____"],
            ["设备工程师", "____", "设备部", "____"],
            ["生产主管", "____", "生产部", "____"],
            ["设计工程师", "____", "研发部", "____"],
            ["SQE（供应商质量工程师）", "____", "质量部", "____"],
            ["质量经理（审核）", "____", "质量部", "____"],
        ],
        col_widths_cm=[5.5, 3.5, 2.5, 4.5],
    )

    # ============ 二、D3 临时遏制措施 ============
    add_heading(doc, "二、D3 临时遏制措施", level=1)

    add_heading(doc, "2.1 遏制措施清单", level=2)
    d3 = template.get("d3_template", {})
    actions = d3.get("containment_actions", [])
    while len(actions) < 5:
        actions.append("____（请补充遏制措施）")

    d3_rows = []
    for i, action in enumerate(actions[:5], start=1):
        d3_rows.append([str(i), action, "____", "____", "____（如100%全检记录/不良率对比）", "待执行"])

    add_table(
        doc,
        headers=["序号", "遏制措施", "责任人", "完成时间", "验证方法", "状态"],
        rows=d3_rows,
        col_widths_cm=[1.2, 6.5, 1.8, 2.0, 4.0, 1.5],
    )

    add_heading(doc, "2.2 遏制有效性验证", level=2)
    add_kv_table(doc, [
        ("遏制前不良率", context.get("defect_rate", "____")),
        ("遏制开始日期", "____"),
        ("遏制后不良率（24h内）", "____"),
        ("遏制后不良率（72h内）", "____"),
        ("遏制结论", "____（达标 / 未达标，是否需要继续遏制）"),
        ("遏制措施截止日期", "____"),
        ("客户是否认可", "____"),
    ])

    # ============ 三、D4 根本原因分析 ============
    add_heading(doc, "三、D4 根本原因分析", level=1)

    add_heading(doc, "3.1 5Why 分析", level=2)
    d4 = template.get("d4_template", {})
    five_why = d4.get("5why_path", {})

    # Use new steps array format, fall back to old format
    steps = five_why.get("steps", [])
    why_rows = []
    root_cause_indices = []

    if steps:
        for idx, step in enumerate(steps):
            level = step.get("level", "____")
            question = step.get("question", "____")
            answer = step.get("answer", "____")
            evidence = step.get("evidence", "____")
            why_rows.append([level, question, answer, evidence])
            if "根因" in str(level):
                root_cause_indices.append(idx)
    else:
        why1 = five_why.get("why1", "____")
        why2_hints = five_why.get("why2_hints", [])
        why3_hints = five_why.get("why3_hints", [])
        why4_hints = five_why.get("why4_hints", [])
        why5_root = five_why.get("why5_root", "____")
        why_rows = [
            ["问题", five_why.get("problem", "____"), "—", "—"],
            ["Why 1", "为什么出现该问题？", why1, "____（数据/观察/文件）"],
            ["Why 2", f"为什么：{why1[:30]}...", "\n".join(f"• {h}" for h in why2_hints) if why2_hints else "____", "____"],
            ["Why 3", "为什么会出现上述原因？", "\n".join(f"• {h}" for h in why3_hints) if why3_hints else "____", "____"],
            ["Why 4", "为什么管理/流程未能预防？", "\n".join(f"• {h}" for h in why4_hints) if why4_hints else "____", "____"],
            ["Why 5（根因）", "为什么流程/规范存在缺陷？", why5_root, "____（流程文件审查）"],
        ]
        root_cause_indices = [5]

    add_table(
        doc,
        headers=["层级", "问题", "答案", "证据"],
        rows=why_rows,
        col_widths_cm=[2.5, 4.5, 6.0, 3.0],
        root_cause_row_indices=root_cause_indices,
    )

    add_heading(doc, "3.2 鱼骨图 6M 排查", level=2)
    six_m = d4.get("6m_analysis", {})

    if isinstance(six_m, list):
        m_rows = []
        for item in six_m:
            m_name = item.get("m", "____")
            finding = item.get("finding", "____")
            judgment = item.get("judgment", "____")
            m_rows.append([m_name, finding, judgment])
        add_table(
            doc,
            headers=["6M 维度", "排查结果", "判定"],
            rows=m_rows,
            col_widths_cm=[3.0, 8.5, 3.5],
        )
    else:
        m_rows = [
            ["Man（人）", six_m.get("man", "____"), "____", "____（是/否）"],
            ["Machine（机）", six_m.get("machine", "____"), "____", "____（是/否）"],
            ["Material（料）", six_m.get("material", "____"), "____", "____（是/否）"],
            ["Method（法）", six_m.get("method", "____"), "____", "____（是/否）"],
            ["Measurement（测）", six_m.get("measurement", "____"), "____", "____（是/否）"],
            ["Environment（环）", six_m.get("environment", "____"), "____", "____（是/否）"],
        ]
        add_table(
            doc,
            headers=["6M 维度", "候选原因", "证据", "是否根因"],
            rows=m_rows,
            col_widths_cm=[2.8, 7.0, 3.5, 2.2],
        )

    add_heading(doc, "3.3 根本原因总结", level=2)
    rc_summary = d4.get("root_cause_summary", [])
    if rc_summary:
        rc_kv = []
        for rc in rc_summary:
            rc_id = rc.get("id", "____")
            rc_desc = rc.get("description", "____")
            rc_type = rc.get("type", "____")
            rc_kv.append((f"{rc_id}（{rc_type}）", rc_desc))
        add_kv_table(doc, rc_kv, label_width_cm=4.5, value_width_cm=12)
    else:
        add_kv_table(doc, [
            ("RC1（直接原因）", "____（参考 5Why 第 1-2 层结论）"),
            ("RC2（管理原因）", "____（参考 5Why 第 3-4 层结论）"),
            ("RC3（系统原因）", "____（参考 5Why 第 5 层结论）"),
        ], label_width_cm=4.5, value_width_cm=12)

    add_heading(doc, "3.4 根本原因验证结论", level=2)
    verify_text = d4.get("verification", "")
    if verify_text:
        add_paragraph(doc, verify_text)
    else:
        add_kv_table(doc, [
            ("验证方法", "____（现场观察 / 数据收集 / 重现试验）"),
            ("验证数据", "____"),
            ("验证结论", "____（已确认 / 待进一步验证）"),
            ("验证人 / 日期", "____"),
        ], label_width_cm=4.5, value_width_cm=12)

    # ============ 四、D5-D6 永久纠正措施 ============
    add_heading(doc, "四、D5-D6 永久纠正措施", level=1)

    add_heading(doc, "4.1 D5 CA 方案评估矩阵", level=2)
    d5_d6 = template.get("d5_d6_template", {})
    permanent_actions = d5_d6.get("permanent_actions", [])

    ca_rows = []
    for i, action in enumerate(permanent_actions, start=1):
        if isinstance(action, dict):
            action_text = action.get("action", "____")
            target = action.get("target", "____")
        else:
            action_text = str(action)
            target = "____"
        ca_rows.append([str(i), action_text, target, "____（高/中/低）", "____（是否引入新风险）", "____（采纳/否决）"])

    while len(ca_rows) < 3:
        ca_rows.append([str(len(ca_rows) + 1), "____（请补充 CA 方案）", "____", "____", "____", "____"])

    add_table(
        doc,
        headers=["序号", "CA 方案", "针对根因", "可行性", "风险评估", "决策"],
        rows=ca_rows,
        col_widths_cm=[1.0, 6.5, 1.8, 1.8, 4.0, 1.8],
    )

    add_heading(doc, "4.2 D6 实施计划跟踪表", level=2)
    d6_rows = []
    for i, action in enumerate(permanent_actions, start=1):
        if isinstance(action, dict):
            action_text = action.get("action", "____")
            target = action.get("target", "____")
            responsible = action.get("responsible", "____")
            due = action.get("due_date", "____")
        else:
            action_text = str(action)
            target = "____"
            responsible = "____"
            due = "____"
        d6_rows.append([str(i), action_text, target, responsible, due, "____", "____"])

    while len(d6_rows) < 3:
        d6_rows.append([str(len(d6_rows) + 1), "____", "____", "____", "____", "____", "____"])

    add_table(
        doc,
        headers=["序号", "实施措施", "目标根因", "责任人", "完成时间", "状态", "验证结果"],
        rows=d6_rows,
        col_widths_cm=[1.0, 5.5, 1.8, 1.8, 1.8, 2.5, 2.6],
    )

    # ============ 五、D7 预防再发生 ============
    add_heading(doc, "五、D7 预防再发生", level=1)

    add_heading(doc, "5.1 横向展开（Yokoten）措施清单", level=2)
    d7 = template.get("d7_template", {})
    yokoten = d7.get("yokoten", [])
    yokoten_rows = []
    for i, item in enumerate(yokoten, start=1):
        yokoten_rows.append([str(i), item, "____（同类产线/产品/客户）", "____", "____", "____"])

    while len(yokoten_rows) < 4:
        yokoten_rows.append([str(len(yokoten_rows) + 1), "____（请补充 Yokoten 措施）", "____", "____", "____", "____"])

    add_table(
        doc,
        headers=["序号", "横向展开措施", "推广范围", "责任人", "完成时间", "状态"],
        rows=yokoten_rows,
        col_widths_cm=[1.0, 6.5, 3.0, 1.8, 1.8, 1.8],
    )

    add_heading(doc, "5.2 PFMEA 更新", level=2)
    add_kv_table(doc, [
        ("本次失效模式", "____"),
        ("原 RPN", "____"),
        ("更新后 RPN", "____"),
        ("PFMEA 更新人 / 日期", "____"),
    ], label_width_cm=4.5, value_width_cm=12)

    # ============ 六、D8 关闭与签名 ============
    add_heading(doc, "六、D8 关闭与签名", level=1)

    add_heading(doc, "6.1 关闭确认", level=2)
    add_kv_table(doc, [
        ("所有 CA 是否实施完成", "____（是/否）"),
        ("所有 CA 是否验证有效", "____（是/否）"),
        ("客户是否认可", "____（是/否，附客户确认邮件）"),
        ("所有文件是否更新", "____（是/否，SOP/PFMEA/培训教材）"),
        ("横向展开是否完成", "____（是/否）"),
        ("关闭日期", "____"),
        ("关闭结论", "____（同意关闭 / 暂不关闭，原因：____）"),
    ], label_width_cm=6.0, value_width_cm=10.5)

    add_heading(doc, "6.2 经验教训", level=2)
    add_kv_table(doc, [
        ("本次问题处理经验", "____"),
        ("可改进之处", "____"),
        ("对其他产品的启示", "____"),
        ("建议改进的管理流程", "____"),
    ], label_width_cm=4.5, value_width_cm=12)

    add_heading(doc, "6.3 签名栏", level=2)
    add_table(
        doc,
        headers=["角色", "姓名", "签名", "日期"],
        rows=[
            ["编制（质量工程师）", "____", "____", "____"],
            ["审核（质量经理）", "____", "____", "____"],
            ["批准（质量总监）", "____", "____", "____"],
            ["客户确认（如需）", "____", "____", "____"],
        ],
        col_widths_cm=[5.0, 4.0, 4.5, 3.0],
    )

    # 末尾报告编号
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    end_p.paragraph_format.space_before = Pt(12)
    end_run = end_p.add_run(f"8D 报告编号：{report_number}")
    set_run_font(end_run, size=9, color="666666")

    # 自动填充模式：遍历所有表格，把 ____ 替换为合理示例值
    if auto_fill:
        _apply_auto_fill_word(doc, report_number)
        print(f"[INFO] Word 已启用自动填充模式，所有 ____ 已替换为示例值")

    doc.save(output_path)
    print(f"[OK] Word 已生成：{output_path}")


# ============================================================
# 主入口
# ============================================================

def parse_args():
    parser = argparse.ArgumentParser(
        description="8D 问题解决报告生成器（Excel + Word）"
    )
    parser.add_argument("--product", required=True, help="产品名称")
    parser.add_argument("--defect", required=True, help="缺陷描述")
    parser.add_argument("--customer", required=False, default="____", help="客户名称")
    parser.add_argument("--defect-rate", required=False, default="____", help="不良率")
    parser.add_argument("--batch-size", required=False, default="____", help="批量数量")
    parser.add_argument(
        "--template",
        required=False,
        default="generic-defect",
        choices=["paint-defect", "assembly-defect", "welding-defect", "dimensional-defect", "generic-defect"],
        help="选用模板 slug",
    )
    parser.add_argument(
        "--output-dir",
        required=False,
        default=os.path.expanduser("~/Desktop"),
        help="输出目录（默认 ~/Desktop）",
    )
    parser.add_argument(
        "--five-why-json",
        required=False,
        default=None,
        help='动态 5Why 内容（JSON 字符串），覆盖模板预填的 5why_path。格式: [{"level":"Why 1","question":"...","answer":"...","evidence":"..."},...]',
    )
    parser.add_argument(
        "--auto-fill",
        action="store_true",
        default=False,
        help="自动填充模式：把所有 ____ 空白替换为合理示例值（化名/示例日期/角色分配）。用户明确说你帮我填时启用。",
    )
    parser.add_argument(
        "--rc-summary-json",
        required=False,
        default=None,
        help='动态根因总结（JSON 字符串），覆盖模板预填的 root_cause_summary。格式: [{"id":"RC1","description":"...","type":"直接原因"},...]',
    )
    parser.add_argument(
        "--containment-actions-json",
        required=False,
        default=None,
        help='动态 D3 遏制措施（JSON 字符串），覆盖模板预填。格式: ["措施1","措施2",...]',
    )
    parser.add_argument(
        "--permanent-actions-json",
        required=False,
        default=None,
        help='动态 D5-D6 永久纠正措施（JSON 字符串），覆盖模板预填。格式: [{"action":"措施描述","target":"针对根因","responsible":"责任人","due_date":"完成时间"},...]',
    )
    parser.add_argument(
        "--yokoten-actions-json",
        required=False,
        default=None,
        help='动态 D7 横向展开措施（JSON 字符串），覆盖模板预填。格式: ["措施1","措施2",...]',
    )
    return parser.parse_args()


def main():
    # Windows 终端 UTF-8 编码修复（避免 print 中文乱码导致 Agent 解析失败）
    import io
    if sys.platform == "win32":
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")

    args = parse_args()

    # 构造 context
    context = {
        "product": args.product,
        "defect": args.defect,
        "customer": args.customer,
        "defect_rate": normalize_defect_rate(args.defect_rate),
        "batch_size": args.batch_size,
    }
    # 行业常识基准警告（>0.5% 视为严重事故）
    if args.defect_rate and args.defect_rate != "____":
        import re
        m = re.search(r"(\d+(?:\.\d+)?)", args.defect_rate)
        if m:
            num = float(m.group(1))
            upper = args.defect_rate.upper()
            is_ppm = "PPM" in upper
            is_percent = "%" in args.defect_rate
            if is_ppm and num > 5000:
                print(f"[WARN] 不良率 {args.defect_rate} > 5000 PPM (0.5%)，属于停发/召回级事故，建议在 D0 阶段标注「严重度=高，建议升级处理」")
            elif is_percent and num > 0.5:
                print(f"[WARN] 不良率 {args.defect_rate} > 0.5%（5000 PPM），属于停发/召回级事故，建议在 D0 阶段标注「严重度=高，建议升级处理」")

    # 生成 8D 编号
    report_number = generate_8d_number()
    print(f"[INFO] 8D 报告编号：{report_number}")

    # 加载模板
    template = load_template(args.template, context)
    print(f"[INFO] 使用模板：{template.get('slug', args.template)} - {template.get('name', '')}")

    # 动态 5Why 覆盖（如果传入了 --five-why-json）
    if args.five_why_json:
        try:
            custom_steps = json.loads(args.five_why_json)
            if isinstance(custom_steps, list) and len(custom_steps) > 0:
                # 确保 d4_template 存在
                if "d4_template" not in template:
                    template["d4_template"] = {}
                # 覆盖 5why_path.steps
                if "5why_path" not in template["d4_template"]:
                    template["d4_template"]["5why_path"] = {"problem": f"产品{args.product}为什么出现{args.defect}？"}
                template["d4_template"]["5why_path"]["steps"] = custom_steps
                print(f"[INFO] 已覆盖 5Why 路径：{len(custom_steps)} 步（动态传入）")
            else:
                print(f"[WARN] --five-why-json 解析后非列表或为空，忽略，使用模板预填 5Why")
        except json.JSONDecodeError as e:
            print(f"[WARN] --five-why-json JSON 解析失败: {e}，使用模板预填 5Why")

    # 动态 RC 总结覆盖（如果传入了 --rc-summary-json）
    if args.rc_summary_json:
        try:
            custom_rc = json.loads(args.rc_summary_json)
            if isinstance(custom_rc, list) and len(custom_rc) > 0:
                # 确保 d4_template 存在
                if "d4_template" not in template:
                    template["d4_template"] = {}
                # 覆盖 root_cause_summary
                template["d4_template"]["root_cause_summary"] = custom_rc
                print(f"[INFO] 已覆盖 RC 总结：{len(custom_rc)} 条（动态传入）")
            else:
                print(f"[WARN] --rc-summary-json 解析后非列表或为空，忽略，使用模板预填 RC")
        except json.JSONDecodeError as e:
            print(f"[WARN] --rc-summary-json JSON 解析失败: {e}，使用模板预填 RC")

    # 动态 D3 遏制措施覆盖（如果传入了 --containment-actions-json）
    if args.containment_actions_json:
        try:
            custom_actions = json.loads(args.containment_actions_json)
            if isinstance(custom_actions, list) and len(custom_actions) > 0:
                if "d3_template" not in template:
                    template["d3_template"] = {}
                template["d3_template"]["containment_actions"] = custom_actions
                print(f"[INFO] 已覆盖 D3 遏制措施：{len(custom_actions)} 条（动态传入）")
            else:
                print(f"[WARN] --containment-actions-json 解析后非列表或为空，忽略")
        except json.JSONDecodeError as e:
            print(f"[WARN] --containment-actions-json JSON 解析失败: {e}，忽略")

    # 动态 D5-D6 永久纠正措施覆盖（如果传入了 --permanent-actions-json）
    if args.permanent_actions_json:
        try:
            custom_pa = json.loads(args.permanent_actions_json)
            if isinstance(custom_pa, list) and len(custom_pa) > 0:
                if "d5_d6_template" not in template:
                    template["d5_d6_template"] = {}
                template["d5_d6_template"]["permanent_actions"] = custom_pa
                print(f"[INFO] 已覆盖 D5-D6 永久纠正措施：{len(custom_pa)} 条（动态传入）")
            else:
                print(f"[WARN] --permanent-actions-json 解析后非列表或为空，忽略")
        except json.JSONDecodeError as e:
            print(f"[WARN] --permanent-actions-json JSON 解析失败: {e}，忽略")

    # 动态 D7 横向展开措施覆盖（如果传入了 --yokoten-actions-json）
    if args.yokoten_actions_json:
        try:
            custom_yk = json.loads(args.yokoten_actions_json)
            if isinstance(custom_yk, list) and len(custom_yk) > 0:
                if "d7_template" not in template:
                    template["d7_template"] = {}
                template["d7_template"]["yokoten"] = custom_yk
                print(f"[INFO] 已覆盖 D7 横向展开措施：{len(custom_yk)} 条（动态传入）")
            else:
                print(f"[WARN] --yokoten-actions-json 解析后非列表或为空，忽略")
        except json.JSONDecodeError as e:
            print(f"[WARN] --yokoten-actions-json JSON 解析失败: {e}，忽略")

    # 确保输出目录存在
    output_dir = os.path.expanduser(args.output_dir)
    os.makedirs(output_dir, exist_ok=True)
    print(f"[INFO] 输出目录：{output_dir}")

    # 生成文件名
    safe_product = safe_filename(args.product)
    safe_defect = safe_filename(args.defect)

    excel_filename = f"8D报告_{safe_product}_{safe_defect}_模板.xlsx"
    word_filename = f"8D报告_{safe_product}_{safe_defect}.docx"

    excel_path = os.path.join(output_dir, excel_filename)
    word_path = os.path.join(output_dir, word_filename)

    # 生成 Excel
    generate_excel(context, template, excel_path, report_number, auto_fill=args.auto_fill)

    # 生成 Word
    generate_word(context, template, word_path, report_number, auto_fill=args.auto_fill)

    # 输出 JSON 结果供调用方解析
    result = {
        "status": "success",
        "report_number": report_number,
        "template_slug": template.get("slug", args.template),
        "template_name": template.get("name", ""),
        "excel_path": excel_path,
        "word_path": word_path,
        "output_dir": output_dir,
    }
    print("\n[RESULT_JSON]")
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
